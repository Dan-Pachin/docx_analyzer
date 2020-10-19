from docx import Document

def test1():
    # This is a very basic testing function, update it as you see fit. (I like to make them colorful and very organized - like comments and beautiful, readable output
    # in the terminal. It takes a little bit of time to set up, but then copy paste takes seconds and you'll be finding failed tests and bugs much faster.)
    test_passed = true
    input_doc = "/Users/Documents/doc1.txt"
    expected_output = "This text should appear!"
    output = run(input_doc)
    
    if(expected_output != output):
        test_passed = false
        
    print(test_passed)

def printResult(pars):
    for index, paragraph in enumerate(output):
        print(paragraph)
    
def run(doc):

    # open selected document
    try:
        this_document = Document(doc)
    except:
        return print('Not a valid docx document!')

    # Name paragraphs and tables lists
    all_paragraphs = this_document.paragraphs
    all_tables = this_document.tables

    # Prompt user for outout type
    mode_selector = input('Show without empty paragraphs? (y/n):')
    while mode_selector != 'y' and mode_selector != 'n':
        mode_selector = input("Please select 'y' or 'n' : ")
    output = []
    # list all paragraphs with their respective text
    for index, paragraph in enumerate(all_paragraphs):
        if paragraph.text:
            output.append(Paragraph[' + str(index) + '] text: ' + str(paragraph.text.strip())
        else:
            if mode_selector == 'y':
                output.append('Paragraph[' + str(index) + '] text: ' + 'This paragraph is empty!')
            else:
                pass
    return output


if __name__ == '__main__':
    output = run(input('Please enter document path: '))
    printResult(output)
    output = run(input('Please enter document path: '))
    printResult(output)
    output = run(input('Please enter document path: '))
    printResult(output)
                          
if __name__ == '__test__':
    test1(input)
    test2(input)
